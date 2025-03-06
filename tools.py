from agno.tools import Toolkit
import os
from docx import Document

class DocumentTools(Toolkit):
    """
    DocumentTools is a toolkit for handling operations on .docx files.
    It provides methods to list, read, and write .docx files in a specified directory.
    """
    def __init__(self):
        """
        Initializes the DocumentTools toolkit and registers its methods.
        """
        super().__init__(name="document_tools")
        self.register(self.list_files)
        self.register(self.read_docx)
        self.register(self.write_docx)

    def list_files(self, folder_path: str) -> str:
        """
        Lists all .docx files in the specified folder and returns them as a comma-separated string.

        Args:
            folder_path (str): The path to the folder where .docx files are to be listed.

        Returns:
            str: A comma-separated string of .docx filenames in the specified folder.
        """
        return ', '.join([f for f in os.listdir(folder_path) if f.endswith('.docx')])

    def read_docx(self, file_path: str) -> str:
        """
        Reads the content of a .docx file and returns it as a string.

        Args:
            file_path (str): The path to the .docx file to be read.

        Returns:
            str: The text content of the .docx file.
        """
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    def write_docx(self, file_path: str, text: str) -> str:
        """
        Writes the provided text to a .docx file and returns the operation status.

        Args:
            file_path (str): The path where the .docx file will be saved.
            text (str): The text content to write into the .docx file.

        Returns:
            str: A message indicating the success or failure of the operation, along with the file name.
        """
        try:
            doc = Document()
            for line in text.split("\n"):
                doc.add_paragraph(line)
            doc.save(file_path)
            return f"Success: File '{file_path}' created."
        except Exception as e:
            return f"Failed: Could not create file '{file_path}'. Error: {e}"